from pydoc import doc
from docx import Document
# docx.shared 用于设置字体大小(图片)等信息
from docx.shared import Cm,Pt

# 创建word文档对象
document = Document()
print(document)

# 添加大标题
document.add_heading("快快乐乐每一天",0)

# 添加段落
p = document.add_paragraph("添加段落的方法")
# 在段落的基础上添加内容 add_run()
run = p.add_run("最新新闻广西百色疫情特别严重,如果是百色的人,注意防护")
run.bold = True # 是否加粗  True表示加粗   False表示不加粗
run.font.size = Pt(18)  # 字体大小设置为18
run.underline = True  # True表示添加下划线  False表示不添加下划线

# 添加一级标题
document.add_heading("一级标题",level=1)
# 添加二级标题
document.add_heading("二级标题",level=2)
# 添加三级标题
document.add_heading("三级标题",level=3)

# 添加带样式的段落
p = document.add_paragraph("带样式的段落",style = "Intense Quote")


# 添加无序列表  style = "List Bullet" 表示无序列表
document.add_paragraph("无序列表第一条数据",style="List Bullet")
document.add_paragraph("无序列表第二条数据",style="List Bullet")
document.add_paragraph("无序列表第三条数据",style="List Bullet")


# 添加有序列表 style = "List Number" 表示有序列表
document.add_paragraph("有序列表第一条数据",style="List Number")
document.add_paragraph("有序列表第二条数据",style="List Number")
document.add_paragraph("有序列表第三条数据",style="List Number")

# 向word文件中添加图片  (注意:图片的路径必须真实存在)
document.add_picture(r"C:\python基础\26-python办公自动化-操作word\resources\meinv.jpg",width=Cm(5.2))

# 向word文件中添加表格
records = (
    ("欣迪","美女","2000-02-11"),
    ("周超","靓仔","2000-03-01"),
    ("文定","猛男","2001-11-14"),
    ("世林","男","1999-04-21"),
    ("俊德","男","1998-03-16")
)

table = document.add_table(rows=1,cols=3)  # rows=1表示从第一行开始
table.style = "Table Grid"  # 表示表格的样式
head_cells = table.rows[0].cells  #添加表格的标题
head_cells[0].text = "姓名"
head_cells[1].text = "性别"
head_cells[2].text = "出生日期"

# 向表格中添加数据
for name,sex,birthday in records:
    row_cells = table.add_row().cells  # 向表格中的每一行添加数据
    row_cells[0].text = name
    row_cells[1].text = sex
    row_cells[2].text = birthday

# 保存文件
document.save("创建word文件.docx")