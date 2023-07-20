from docx import Document
doc = Document(r"C:\python基础\26-python办公自动化-操作word\resources\离职证明模板.docx")

for no,p in enumerate(doc.paragraphs):  # no表示行号   p表示段落   doc.paragraphs表示文档的所有的段落
    print(no,p.text)   # p.text表示段落的内容